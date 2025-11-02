
import asyncio
import logging
import os
import tempfile
import subprocess
import io
import yt_dlp
import httpx
import uuid
import json
import requests
import time
import secrets
from typing import List, Dict, Optional, Callable, Any, Tuple, TypedDict

class Segment(TypedDict):
    speaker: str
    text: str
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image, ImageDraw, ImageFont

from .config import (
    ASSEMBLYAI_BASE_URL, HEADERS, API_TIMEOUT, FFMPEG_BIN, FFPROBE_BIN,
    SEGMENT_DURATION, OPENROUTER_API_KEYS, OPENROUTER_BASE_URL, OPENROUTER_MODEL, FONT_PATH,
    YOOMONEY_WALLET, YOOMONEY_BASE_URL, SUBSCRIPTION_AMOUNT, THUMBNAIL_COLOR
)
from .exceptions import PaymentError, TranscriptionError, FileProcessingError, APIError, NetworkError
from .circuit_breaker import CircuitBreaker


logger = logging.getLogger(__name__)


# =============================
#     OpenRouter Client with API Key Rotation
# =============================
class OpenRouterClient:
    """Клиент для работы с OpenRouter API с автоматической ротацией ключей."""

    def __init__(self, base_url: str = OPENROUTER_BASE_URL, model: str = OPENROUTER_MODEL):
        self.base_url = base_url
        self.model = model
        logger.info("Инициализирован OpenRouter клиент с API key manager")

    async def make_request(self, messages: List[Dict[str, str]], temperature: float = 0.2) -> str:
        """Выполнить запрос к OpenRouter с автоматической ротацией ключей через API key manager."""
        from .services.security import api_key_manager

        url = f"{self.base_url}/chat/completions"
        data = {
            "model": self.model,
            "messages": messages,
            "temperature": temperature
        }

        # Try with different keys if available
        max_attempts = len(api_key_manager._keys) if api_key_manager._keys else 1

        for attempt in range(max_attempts):
            api_key = api_key_manager.get_current_key()
            if not api_key:
                logger.error("OPENROUTER_API_KEYS не настроены")
                raise ValueError("OPENROUTER_API_KEYS не настроены")

            logger.debug(f"Попытка запроса к OpenRouter с ключом (attempt {attempt + 1}/{max_attempts})")
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }

            try:
                async with httpx.AsyncClient() as client:
                    response = await client.post(url, headers=headers, json=data, timeout=60)
                    if response.status_code == 429:
                        logger.warning("Получен 429 (Too Many Requests), пробуем следующий ключ")
                        api_key_manager._rotate_key()
                        continue
                    response.raise_for_status()

                    # Mark key as used for rotation tracking
                    api_key_manager.mark_key_used(api_key)

                    return response.json()['choices'][0]['message']['content'].strip()
            except httpx.HTTPStatusError as e:
                if e.response.status_code == 429:
                    logger.warning("HTTP 429, пробуем следующий ключ")
                    api_key_manager._rotate_key()
                    continue
                else:
                    logger.error(f"OpenRouter API ошибка {e.response.status_code}: {e}")
                    raise
            except Exception as e:
                logger.error(f"Ошибка запроса к OpenRouter: {e}")
                if attempt < max_attempts - 1:
                    api_key_manager._rotate_key()
                    continue
                raise

        logger.error("Все OPENROUTER API ключи вернули ошибки или 429")
        raise Exception("Все OPENROUTER API ключи исчерпаны")


# Создаём экземпляр клиента
openrouter_client = OpenRouterClient()


# =============================
#     YooMoney Payment
# =============================
async def create_yoomoney_payment(user_id: int, amount: int, description: str) -> Tuple[Optional[str], Optional[str]]:
    """Создает ссылку на оплату YooMoney."""
    payment_label = f"sub_{user_id}_{uuid.uuid4()}"
    quickpay_url = f"{YOOMONEY_BASE_URL}/quickpay/confirm.xml"
    params = {
        "receiver": YOOMONEY_WALLET,
        "quickpay-form": "shop",
        "targets": description,
        "paymentType": "SB",
        "sum": amount,
        "label": payment_label,
    }

    async def _make_request():
        async with httpx.AsyncClient() as client:
            response = await client.post(quickpay_url, data=params, follow_redirects=False)
            if response.status_code == 302:
                # Для YooMoney редирект 302 - это успешный ответ
                redirect_url = response.headers.get('Location', '')
                if redirect_url:
                    return redirect_url, payment_label
                else:
                    raise PaymentError("YooMoney не вернул URL для оплаты")
            else:
                response.raise_for_status()
                from urllib.parse import urlencode
                encoded_params = urlencode(params)
                payment_url = f"{YOOMONEY_BASE_URL}/quickpay/confirm.xml?{encoded_params}"
                return payment_url, payment_label

    circuit_breaker = CircuitBreaker(failure_threshold=3, recovery_timeout=30, expected_exception=(httpx.RequestError,))

    for attempt in range(3):
        try:
            result = await circuit_breaker.call(_make_request)
            logger.info(f"Создана ссылка на оплату для user_id {user_id}: {payment_label}")
            return result
        except (httpx.RequestError,) as e:
            logger.warning(f"Попытка {attempt + 1}/3 создания платежа YooMoney не удалась: {e}")
            if attempt == 2:
                raise PaymentError(f"Не удалось создать платеж YooMoney: {e}") from e
            time.sleep(2 ** attempt)
    return None, None


# =============================
#     Регистрация шрифта PDF
# =============================
try:
    pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
except Exception as e:
    logger.error(f"Failed to register DejaVu: {e}")
    # ничего не регистрируем: Helvetica встроенная


# ---------- Сохранение в разные форматы ----------

def _register_pdf_font_if_needed() -> None:
    try:
        if 'DejaVu' not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
    except Exception:
        pass

def save_text_to_pdf(text: str, output_path: str) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib.pagesizes import A4

    # Ensure text is properly encoded as UTF-8
    if isinstance(text, str):
        text = text.encode('utf-8').decode('utf-8')

    # Create PDF with canvas for better UTF-8 support
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    # Set margins
    left_margin = inch
    right_margin = inch
    top_margin = inch
    bottom_margin = inch
    available_width = width - left_margin - right_margin

    # Try to register and use DejaVu font
    try:
        if 'DejaVu' not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
        font_name = "DejaVu"
        c.setFont(font_name, 12)
    except Exception as e:
        logger.warning(f"Could not use DejaVu font: {e}, falling back to Helvetica")
        font_name = "Helvetica"
        c.setFont(font_name, 12)

    # Get font metrics for text wrapping
    font_size = 12
    line_height = 14

    # Function to wrap text to fit within available width
    def wrap_text(text_line, max_width):
        """Wrap text to fit within max_width, returning list of lines"""
        if not text_line.strip():
            return [text_line]

        words = text_line.split()
        lines = []
        current_line = ""

        for word in words:
            # Check if adding this word would exceed the width
            test_line = current_line + " " + word if current_line else word
            if pdfmetrics.stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                # If current_line is not empty, add it to lines
                if current_line:
                    lines.append(current_line)
                # Start new line with current word
                if pdfmetrics.stringWidth(word, font_name, font_size) <= max_width:
                    current_line = word
                else:
                    # Word itself is too long, force break it
                    current_line = word[:int(len(word) * (max_width / pdfmetrics.stringWidth(word, font_name, font_size)))]
                    lines.append(current_line)
                    current_line = word[len(current_line):]

        # Add remaining line
        if current_line:
            lines.append(current_line)

        return lines if lines else [""]

    # Split text into paragraphs and process each
    paragraphs = text.split('\n\n')
    y_position = height - top_margin

    for paragraph in paragraphs:
        # Skip empty paragraphs
        if not paragraph.strip():
            continue

        # Split paragraph into lines
        lines = paragraph.split('\n')

        for line in lines:
            # Wrap the line
            wrapped_lines = wrap_text(line, available_width)

            for wrapped_line in wrapped_lines:
                # Check if we need a new page
                if y_position < bottom_margin + line_height:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y_position = height - top_margin

                # Draw the line
                c.drawString(left_margin, y_position, wrapped_line)
                y_position -= line_height

        # Add extra space between paragraphs
        y_position -= line_height * 0.5

    c.save()


def save_text_to_txt(text: str, output_path: str) -> None:
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def save_text_to_md(text: str, output_path: str) -> None:
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def save_text_to_docx(text: str, output_path: str) -> None:
    try:
        from docx import Document
        doc = Document()
        for par in text.split("\n\n"):
            for line in par.split("\n"):
                doc.add_paragraph(line)
            doc.add_paragraph("")
        doc.save(output_path)
    except Exception as e:
        logger.warning(f"Не удалось сохранить DOCX ({e}), сохраняю как TXT")
        save_text_to_txt(text, output_path)


# ---------- Аудио-обработка / API ----------
class AudioProcessor:
    @staticmethod
    def split_audio(input_path: str, segment_time: int = SEGMENT_DURATION) -> list[str]:
        output_dir = tempfile.mkdtemp(prefix="fragments_")
        output_pattern = os.path.join(output_dir, "fragment_%03d.mp3")
        ffmpeg_path = FFMPEG_BIN
        command = [ffmpeg_path, "-i", input_path, "-f", "segment", "-segment_time", str(segment_time), "-c", "copy", output_pattern]

        try:
            subprocess.run(command, check=True, capture_output=True, text=True)
            return sorted([
                os.path.join(output_dir, f)
                for f in os.listdir(output_dir)
                if f.startswith("fragment_") and f.endswith(".mp3")
            ])
        except subprocess.CalledProcessError as e:
            logger.error(f"FFmpeg error: {e.stderr}")
            raise RuntimeError("Ошибка при разделении аудио") from e

    @staticmethod
    def cleanup(files: List[str]) -> None:
        for path in files:
            try:
                if os.path.isfile(path):
                    os.remove(path)
                elif os.path.isdir(path):
                    for f in os.listdir(path):
                        os.remove(os.path.join(path, f))
                    os.rmdir(path)
            except (OSError, FileNotFoundError) as e:
                logger.warning(f"Ошибка удаления {path}: {e}")


async def upload_to_assemblyai(file_path: str, retries: int = 3) -> str:
    async def _make_request():
        async with httpx.AsyncClient() as client:
            with open(file_path, "rb") as f:
                response = await client.post(
                    f"{ASSEMBLYAI_BASE_URL}/upload",
                    headers=HEADERS,
                    files={"file": f},
                    timeout=API_TIMEOUT
                )
            response.raise_for_status()
            return response.json()["upload_url"]

    circuit_breaker = CircuitBreaker(failure_threshold=3, recovery_timeout=30, expected_exception=(httpx.RequestError, httpx.HTTPStatusError, KeyError))

    for attempt in range(retries):
        try:
            result = await circuit_breaker.call(_make_request)
            return result
        except (httpx.RequestError, httpx.HTTPStatusError, KeyError) as e:
            logger.warning(f"Попытка {attempt + 1}/{retries} загрузки файла не удалась: {str(e)}")
            if attempt == retries - 1:
                raise TranscriptionError("Не удалось загрузить файл на сервер AssemblyAI") from e
            time.sleep(2 ** attempt)


async def transcribe_with_assemblyai(audio_url: str, retries: int = 3) -> Dict[str, Any]:
    headers = {
        "authorization": HEADERS['authorization'],
        "content-type": "application/json"
    }
    payload = {
        "audio_url": audio_url,
        "speaker_labels": True,
        "punctuate": True,
        "format_text": True,
        "language_detection": True
    }

    async def _make_request():
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                f"{ASSEMBLYAI_BASE_URL}/transcript",
                headers=headers, json=payload
            )
            resp.raise_for_status()
            transcript_id = resp.json()["id"]
            while True:
                status = await client.get(
                    f"{ASSEMBLYAI_BASE_URL}/transcript/{transcript_id}",
                    headers=headers
                )
                result = status.json()
                if result["status"] == "completed":
                    return result
                elif result["status"] == "error":
                    raise TranscriptionError(result["error"])
                await asyncio.sleep(3)

    circuit_breaker = CircuitBreaker(failure_threshold=3, recovery_timeout=30, expected_exception=(httpx.RequestError, httpx.HTTPStatusError, TranscriptionError, KeyError))

    for attempt in range(retries):
        try:
            result = await circuit_breaker.call(_make_request)
            return result
        except (httpx.RequestError, httpx.HTTPStatusError, TranscriptionError, KeyError) as e:
            logger.warning(f"Попытка {attempt + 1}/{retries} транскрипции не удалась: {str(e)}")
            if attempt == retries - 1:
                raise TranscriptionError("Не удалось выполнить транскрипцию") from e
            time.sleep(2 ** attempt)


async def download_youtube_audio(url: str, progress_callback: Optional[Callable] = None) -> str:
    loop = asyncio.get_running_loop()
    progress_queue = asyncio.Queue()

    def progress_hook(data):
        if data['status'] == 'downloading' and progress_callback:
            try:
                percent_str = data.get('_percent_str', '0%')
                percent_value = float(percent_str.strip().replace('%', ''))
                loop.call_soon_threadsafe(progress_queue.put_nowait, percent_value)
            except:
                pass

    def sync_download():
        temp_dir = tempfile.gettempdir()
        unique_id = str(uuid.uuid4())
        outtmpl = os.path.join(temp_dir, f"{unique_id}")
        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": outtmpl,
            "ffmpeg_location": os.path.dirname(FFMPEG_BIN) or None,

            "progress_hooks": [progress_hook],
            "postprocessors": [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
            }],
            "quiet": True,
            "no_warnings": False,
            "extract_flat": False,
            "http_headers": {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            },
        }
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
                expected_filename = f"{outtmpl}.mp3"
                if os.path.exists(expected_filename):
                    return expected_filename
                else:
                    for file in os.listdir(temp_dir):
                        if file.startswith(os.path.basename(unique_id)):
                            return os.path.join(temp_dir, file)
                    raise FileNotFoundError(f"Скачанный аудиофайл не найден: {expected_filename}")
        except Exception as e:
            logger.error(f"Ошибка скачивания YouTube: {str(e)}")
            raise RuntimeError(f"Ошибка скачивания видео: {str(e)}") from e

    async def process_progress():
        while True:
            try:
                data = await asyncio.wait_for(progress_queue.get(), timeout=1.0)
                if progress_callback:
                    await progress_callback(data)
            except asyncio.TimeoutError:
                if download_task.done():
                    break
            except Exception as e:
                logger.warning(f"Ошибка обработки прогресса: {str(e)}")
                break

    download_task = loop.run_in_executor(None, sync_download)
    progress_task = asyncio.create_task(process_progress())
    try:
        result = await download_task
        return result
    finally:
        progress_task.cancel()
        try:
            await progress_task
        except asyncio.CancelledError:
            pass


def format_results_with_speakers(segments: List[Segment]) -> str:
    return "\n\n".join(f"Спикер {seg['speaker']}:\n{seg['text']}" for seg in segments)


def format_results_plain(segments: List[Segment]) -> str:
    return "\n\n".join(seg["text"] for seg in segments)


async def generate_summary_timecodes(segments: List[Segment]) -> str:
    full_text_with_timestamps = ""
    for i, seg in enumerate(segments):
        start_minute = i * SEGMENT_DURATION // 60
        start_second = i * SEGMENT_DURATION % 60
        start_code = f"{start_minute:02}:{start_second:02}"
        full_text_with_timestamps += f"[{start_code}] {seg['text']}\n\n"
    prompt = f"""
Проанализируй полную расшифровку аудио с тайм-кодами и создай структурированное оглавление.
Текст с тайм-кодами:
{full_text_with_timestamps}
Инструкции:
1. Выдели ОСНОВНЫЕ смысловые блоки и темы
2. Группируй несколько последовательных сегментов в один логический блок
3. Для каждого блока укажи время начала
4. Дай емкое описание содержания блока
5. Сохраняй хронологический порядок
Формат ответа:
Тайм-коды
MM:SS - [Основная тема/событие]
[Дополнительные детали]
MM:SS - [Следующая основная тема]
...
"""

    try:
        return await openrouter_client.make_request([{"role": "user", "content": prompt}], temperature=0.2)
    except Exception as e:
        logger.warning(f"Попытка генерации тайм-кодов с OpenRouter не удалась: {e}")
        logger.info("Используем fallback для тайм-кодов")

    # Fallback
    fallback_result = "Тайм-коды\n\n"
    for i, seg in enumerate(segments):
        start_minute = i * SEGMENT_DURATION // 60
        start_second = i * SEGMENT_DURATION % 60
        start_code = f"{start_minute:02}:{start_second:02}"
        fallback_result += f"{start_code} - {seg['text'][:50]}...\n"
    return fallback_result


async def convert_to_mp3(input_path: str) -> str:
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3").name
    ffmpeg_path = FFMPEG_BIN
    command = [ffmpeg_path, "-i", input_path, "-acodec", "libmp3lame", "-q:a", "2", "-y", output_path]

    try:
        process = await asyncio.create_subprocess_exec(
            *command,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        if process.returncode != 0:
            logger.error(f"Ошибка конвертации: {stderr.decode()}")
            raise RuntimeError(f"Ошибка конвертации файла: {stderr.decode()}")
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            logger.error(f"Выходной файл {output_path} не создан или пуст")
            raise RuntimeError("Конвертация не удалась: выходной файл не создан")
        logger.info(f"Конвертация успешна: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Ошибка конвертации {input_path}: {str(e)}")
        raise RuntimeError(f"Ошибка конвертации: {str(e)}") from e


async def process_audio_file(file_path: str, user_id: int, progress_callback: Optional[Callable] = None) -> List[Segment]:
    try:
        logger.info(f"Обработка аудиофайла: {file_path}")
        if progress_callback:
            await progress_callback(0.01, "Загружаю файл для обработки...")
        audio_url = await upload_to_assemblyai(file_path)
        if progress_callback:
            await progress_callback(0.30, "Запускаю транскрибацию...")
        result = await transcribe_with_assemblyai(audio_url)
        if progress_callback:
            await progress_callback(0.90, "Формирую результаты...")
        segments = []
        if "utterances" in result and result["utterances"]:
            for utt in result["utterances"]:
                segments.append({
                    "speaker": utt.get("speaker", "?"),
                    "text": (utt.get("text") or "").strip()
                })
        elif "text" in result:
            segments.append({"speaker": "?", "text": (result["text"] or "").strip()})
        if progress_callback:
            await progress_callback(1.0, "Обработка завершена!")
        logger.info(f"Транскрибация завершена, найдено {len(segments)} сегментов")
        return segments
    except (TranscriptionError, FileProcessingError) as e:
        logger.error(f"Ошибка в process_audio_file: {str(e)}")
        raise


THUMBNAIL_CACHE = {}

def create_custom_thumbnail(thumbnail_path: Optional[str] = None) -> Optional[io.BytesIO]:
    cache_key = thumbnail_path or "default"
    if cache_key in THUMBNAIL_CACHE:
        thumbnail_bytes = io.BytesIO(THUMBNAIL_CACHE[cache_key])
        thumbnail_bytes.seek(0)
        return thumbnail_bytes
    try:
        if thumbnail_path and os.path.exists(thumbnail_path):
            with Image.open(thumbnail_path) as img:
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                target_size = (320, 320)
                img.thumbnail(target_size, Image.LANCZOS)
                square_img = Image.new('RGB', target_size, (255, 255, 255))
                x_offset = (target_size[0] - img.width) // 2
                y_offset = (target_size[1] - img.height) // 2
                square_img.paste(img, (x_offset, y_offset))
                thumbnail_bytes = io.BytesIO()
                square_img.save(thumbnail_bytes, format='JPEG', quality=95, optimize=True)
                thumbnail_bytes.seek(0)
                THUMBNAIL_CACHE[cache_key] = thumbnail_bytes.getvalue()
                thumbnail_bytes.seek(0)
                return thumbnail_bytes
        else:
            target_size = (320, 320)
            img = Image.new('RGB', target_size, color=THUMBNAIL_COLOR)
            draw = ImageDraw.Draw(img)
            margin = 10
            draw.rectangle([margin, margin, target_size[0] - margin, target_size[1] - margin],
                           outline=(255, 255, 255), width=4)
            try:
                font = ImageFont.truetype("arial.ttf", 80)
            except:
                font = ImageFont.load_default()
            text = "PDF"
            bbox = draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            x = (target_size[0] - text_width) // 2
            y = (target_size[1] - text_height) // 2
            draw.text((x, y), text, fill=(255, 255, 255), font=font)
            thumbnail_bytes = io.BytesIO()
            img.save(thumbnail_bytes, format='JPEG', quality=95, optimize=True)
            thumbnail_bytes.seek(0)
            THUMBNAIL_CACHE[cache_key] = thumbnail_bytes.getvalue()
            thumbnail_bytes.seek(0)
            return thumbnail_bytes
    except (IOError, OSError) as e:
        logger.error(f"Ошибка создания thumbnail: {e}")
        return None
