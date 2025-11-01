import logging
import os
import tempfile
import subprocess
import io
import yt_dlp
import asyncio
import uuid
from typing import List, Dict, Optional, Callable, Any, Tuple
from PIL import Image, ImageDraw, ImageFont

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

from ..config import (
    FFMPEG_BIN, FFPROBE_BIN, FONT_PATH, THUMBNAIL_COLOR, CUSTOM_THUMBNAIL_PATH
)
from ..exceptions import FileProcessingError

logger = logging.getLogger(__name__)


# =============================
#     Регистрация шрифта PDF
# =============================
try:
    pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
except Exception as e:
    logger.error(f"Failed to register DejaVu: {e}")
    # ничего не регистрируем: Helvetica встроенная


# ---------- Сохранение в разные форматы ----------

def _register_pdf_font_if_needed() -> None:
    try:
        if 'DejaVu' not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
    except Exception:
        pass

def save_text_to_pdf(text: str, output_path: str) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib.pagesizes import A4

    # Ensure text is properly encoded as UTF-8
    if isinstance(text, str):
        text = text.encode('utf-8').decode('utf-8')

    # Create PDF with canvas for better UTF-8 support
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    # Set margins
    left_margin = inch
    right_margin = inch
    top_margin = inch
    bottom_margin = inch
    available_width = width - left_margin - right_margin

    # Try to register and use DejaVu font
    try:
        if 'DejaVu' not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
        font_name = "DejaVu"
        c.setFont(font_name, 12)
    except Exception as e:
        logger.warning(f"Could not use DejaVu font: {e}, falling back to Helvetica")
        font_name = "Helvetica"
        c.setFont(font_name, 12)

    # Get font metrics for text wrapping
    font_size = 12
    line_height = 14

    # Function to wrap text to fit within available width
    def wrap_text(text_line, max_width):
        """Wrap text to fit within max_width, returning list of lines"""
        if not text_line.strip():
            return [text_line]

        words = text_line.split()
        lines = []
        current_line = ""

        for word in words:
            # Check if adding this word would exceed the width
            test_line = current_line + " " + word if current_line else word
            if pdfmetrics.stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                # If current_line is not empty, add it to lines
                if current_line:
                    lines.append(current_line)
                # Start new line with current word
                if pdfmetrics.stringWidth(word, font_name, font_size) <= max_width:
                    current_line = word
                else:
                    # Word itself is too long, force break it
                    current_line = word[:int(len(word) * (max_width / pdfmetrics.stringWidth(word, font_name, font_size)))]
                    lines.append(current_line)
                    current_line = word[len(current_line):]

        # Add remaining line
        if current_line:
            lines.append(current_line)

        return lines if lines else [""]

    # Split text into paragraphs and process each
    paragraphs = text.split('\n\n')
    y_position = height - top_margin

    for paragraph in paragraphs:
        # Skip empty paragraphs
        if not paragraph.strip():
            continue

        # Split paragraph into lines
        lines = paragraph.split('\n')

        for line in lines:
            # Wrap the line
            wrapped_lines = wrap_text(line, available_width)

            for wrapped_line in wrapped_lines:
                # Check if we need a new page
                if y_position < bottom_margin + line_height:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y_position = height - top_margin

                # Draw the line
                c.drawString(left_margin, y_position, wrapped_line)
                y_position -= line_height

        # Add extra space between paragraphs
        y_position -= line_height * 0.5

    c.save()


def save_text_to_txt(text: str, output_path: str) -> None:
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def save_text_to_md(text: str, output_path: str) -> None:
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def save_text_to_docx(text: str, output_path: str) -> None:
    try:
        from docx import Document
        doc = Document()
        for par in text.split("\n\n"):
            for line in par.split("\n"):
                doc.add_paragraph(line)
            doc.add_paragraph("")
        doc.save(output_path)
    except Exception as e:
        logger.warning(f"Не удалось сохранить DOCX ({e}), сохраняю как TXT")
        save_text_to_txt(text, output_path)


async def download_youtube_audio(url: str, progress_callback: Optional[Callable] = None) -> str:
    loop = asyncio.get_running_loop()
    progress_queue = asyncio.Queue()

    def progress_hook(data):
        if data['status'] == 'downloading' and progress_callback:
            try:
                percent_str = data.get('_percent_str', '0%')
                percent_value = float(percent_str.strip().replace('%', ''))
                loop.call_soon_threadsafe(progress_queue.put_nowait, percent_value)
            except:
                pass

    def sync_download():
        temp_dir = tempfile.gettempdir()
        unique_id = str(uuid.uuid4())
        outtmpl = os.path.join(temp_dir, f"{unique_id}")
        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": outtmpl,
            "ffmpeg_location": os.path.dirname(FFMPEG_BIN) or None,

            "progress_hooks": [progress_hook],
            "postprocessors": [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
            }],
            "quiet": True,
            "no_warnings": False,
            "extract_flat": False,
            "http_headers": {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            },
        }
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
                expected_filename = f"{outtmpl}.mp3"
                if os.path.exists(expected_filename):
                    return expected_filename
                else:
                    for file in os.listdir(temp_dir):
                        if file.startswith(os.path.basename(unique_id)):
                            return os.path.join(temp_dir, file)
                    raise FileNotFoundError(f"Скачанный аудиофайл не найден: {expected_filename}")
        except Exception as e:
            logger.error(f"Ошибка скачивания YouTube: {str(e)}")
            raise RuntimeError(f"Ошибка скачивания видео: {str(e)}") from e

    async def process_progress():
        while True:
            try:
                data = await asyncio.wait_for(progress_queue.get(), timeout=1.0)
                if progress_callback:
                    await progress_callback(data)
            except asyncio.TimeoutError:
                if download_task.done():
                    break
            except Exception as e:
                logger.warning(f"Ошибка обработки прогресса: {str(e)}")
                break

    download_task = loop.run_in_executor(None, sync_download)
    progress_task = asyncio.create_task(process_progress())
    try:
        result = await download_task
        return result
    finally:
        progress_task.cancel()
        try:
            await progress_task
        except asyncio.CancelledError:
            pass


async def convert_to_mp3(input_path: str) -> str:
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3").name
    ffmpeg_path = FFMPEG_BIN
    command = [ffmpeg_path, "-i", input_path, "-acodec", "libmp3lame", "-q:a", "2", "-y", output_path]

    try:
        process = await asyncio.create_subprocess_exec(
            *command,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        if process.returncode != 0:
            logger.error(f"Ошибка конвертации: {stderr.decode()}")
            raise RuntimeError(f"Ошибка конвертации файла: {stderr.decode()}")
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            logger.error(f"Выходной файл {output_path} не создан или пуст")
            raise RuntimeError("Конвертация не удалась: выходной файл не создан")
        logger.info(f"Конвертация успешна: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Ошибка конвертации {input_path}: {str(e)}")
        raise RuntimeError(f"Ошибка конвертации: {str(e)}") from e


THUMBNAIL_CACHE = {}

def create_custom_thumbnail(thumbnail_path: Optional[str] = None) -> Optional[io.BytesIO]:
    cache_key = thumbnail_path or "default"
    if cache_key in THUMBNAIL_CACHE:
        thumbnail_bytes = io.BytesIO(THUMBNAIL_CACHE[cache_key])
        thumbnail_bytes.seek(0)
        return thumbnail_bytes
    try:
        if thumbnail_path and os.path.exists(thumbnail_path):
            with Image.open(thumbnail_path) as img:
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                target_size = (320, 320)
                img.thumbnail(target_size, Image.LANCZOS)
                square_img = Image.new('RGB', target_size, (255, 255, 255))
                x_offset = (target_size[0] - img.width) // 2
                y_offset = (target_size[1] - img.height) // 2
                square_img.paste(img, (x_offset, y_offset))
                thumbnail_bytes = io.BytesIO()
                square_img.save(thumbnail_bytes, format='JPEG', quality=95, optimize=True)
                thumbnail_bytes.seek(0)
                THUMBNAIL_CACHE[cache_key] = thumbnail_bytes.getvalue()
                thumbnail_bytes.seek(0)
                return thumbnail_bytes
        else:
            target_size = (320, 320)
            img = Image.new('RGB', target_size, color=THUMBNAIL_COLOR)
            draw = ImageDraw.Draw(img)
            margin = 10
            draw.rectangle([margin, margin, target_size[0] - margin, target_size[1] - margin],
                           outline=(255, 255, 255), width=4)
            try:
                font = ImageFont.truetype("arial.ttf", 80)
            except:
                font = ImageFont.load_default()
            text = "PDF"
            bbox = draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            x = (target_size[0] - text_width) // 2
            y = (target_size[1] - text_height) // 2
            draw.text((x, y), text, fill=(255, 255, 255), font=font)
            thumbnail_bytes = io.BytesIO()
            img.save(thumbnail_bytes, format='JPEG', quality=95, optimize=True)
            thumbnail_bytes.seek(0)
            THUMBNAIL_CACHE[cache_key] = thumbnail_bytes.getvalue()
            thumbnail_bytes.seek(0)
            return thumbnail_bytes
    except (IOError, OSError) as e:
        logger.error(f"Ошибка создания thumbnail: {e}")
        return None
